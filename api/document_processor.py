import os
from pathlib import Path
from typing import Dict, Any
import PyPDF2
import docx
from langchain.schema import Document

class DocumentProcessor:
    """Process and extract text from various document formats"""
    
    def __init__(self):
        self.supported_extensions = {'.pdf', '.docx', '.doc', '.txt'}
    
    def extract_text(self, file_path: str) -> str:
        """Extract text from a document file"""
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        extension = file_path.suffix.lower()
        
        if extension == '.pdf':
            return self._extract_from_pdf(file_path)
        elif extension in ['.docx', '.doc']:
            return self._extract_from_docx(file_path)
        elif extension == '.txt':
            return self._extract_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {extension}")
    
    def _extract_from_pdf(self, file_path: Path) -> str:
        """Extract text from PDF files"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text.strip()
        except Exception as e:
            raise Exception(f"Error reading PDF: {str(e)}")
    
    def _extract_from_docx(self, file_path: Path) -> str:
        """Extract text from DOCX files"""
        try:
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text.strip()
        except Exception as e:
            raise Exception(f"Error reading DOCX: {str(e)}")
    
    def _extract_from_txt(self, file_path: Path) -> str:
        """Extract text from TXT files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read().strip()
        except Exception as e:
            raise Exception(f"Error reading TXT: {str(e)}")
    
    def create_document(self, file_path: str, doc_type: str) -> Document:
        """Create a LangChain Document object"""
        text = self.extract_text(file_path)
        
        metadata = {
            "source": file_path,
            "type": doc_type,  # "resume" or "job_description"
            "filename": Path(file_path).name
        }
        
        return Document(page_content=text, metadata=metadata)
    
    def process_resume(self, file_path: str) -> Dict[str, Any]:
        """Process a resume and extract structured information"""
        document = self.create_document(file_path, "resume")
        
        # Basic processing - we'll enhance this with LLM analysis later
        return {
            "raw_text": document.page_content,
            "metadata": document.metadata,
            "word_count": len(document.page_content.split()),
            "char_count": len(document.page_content)
        }
    
    def process_job_description(self, file_path: str) -> Dict[str, Any]:
        """Process a job description and extract structured information"""
        document = self.create_document(file_path, "job_description")
        
        # Basic processing - we'll enhance this with LLM analysis later
        return {
            "raw_text": document.page_content,
            "metadata": document.metadata,
            "word_count": len(document.page_content.split()),
            "char_count": len(document.page_content)
        }